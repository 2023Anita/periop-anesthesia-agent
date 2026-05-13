from __future__ import annotations

from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path


@dataclass(frozen=True)
class ExtractedDocument:
    text: str
    notes: list[str] = field(default_factory=list)


def extract_document_text(filename: str, content: bytes) -> ExtractedDocument:
    suffix = Path(filename).suffix.lower()
    if suffix in {".txt", ".md"}:
        return ExtractedDocument(_decode_text(content), ["已按纯文本资料抽取。"])
    if suffix == ".pdf":
        return _extract_pdf(content)
    if suffix in {".docx"}:
        return _extract_docx(content)
    if suffix in {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff"}:
        return _extract_image_ocr(content)
    return ExtractedDocument(
        _decode_text(content),
        [f"未知文件类型 {suffix or '无后缀'}，已尝试按文本读取。"],
    )


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "gb18030", "latin-1"):
        try:
            return content.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return ""


def _extract_pdf(content: bytes) -> ExtractedDocument:
    try:
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(content))
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
        text = "\n\n".join(page for page in pages if page)
        notes = [f"已抽取 PDF 文本，共 {len(reader.pages)} 页。"]
        if not text:
            notes.append("PDF 未抽取到文本，可能是扫描件；建议改用 OCR 图片或补充文字。")
        return ExtractedDocument(text, notes)
    except Exception as exc:  # pragma: no cover - defensive dependency boundary
        return ExtractedDocument("", [f"PDF 抽取失败：{exc}"])


def _extract_docx(content: bytes) -> ExtractedDocument:
    try:
        from docx import Document

        document = Document(BytesIO(content))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
        return ExtractedDocument(text, ["已抽取 Word 文档正文。"])
    except Exception as exc:  # pragma: no cover - defensive dependency boundary
        return ExtractedDocument("", [f"Word 抽取失败：{exc}"])


def _extract_image_ocr(content: bytes) -> ExtractedDocument:
    try:
        from PIL import Image
        import pytesseract

        image = Image.open(BytesIO(content))
        text = pytesseract.image_to_string(image, lang="chi_sim+eng").strip()
        notes = ["已尝试对图片执行 OCR。"]
        if not text:
            notes.append("图片 OCR 未得到有效文字；建议上传清晰报告或手动粘贴心电图结论。")
        return ExtractedDocument(text, notes)
    except Exception as exc:
        return ExtractedDocument(
            "",
            [
                "图片 OCR 接口已预留，但本机 OCR 环境不可用或图片无法识别。",
                f"OCR 失败原因：{exc}",
            ],
        )

